from docx import Document
from docx.shared import Inches

document = Document('example.docx')

for paragraph in document.paragraphs:
    if "[image]" in paragraph.text:
        paragraph.text = paragraph.text.strip().replace("[image]", "")
        run = paragraph.add_run()
        run.add_picture('FE_logo.jpg', width=Inches(3))
        break


document.save('example.docx')